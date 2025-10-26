"""
Script to generate a formatted DOCX report from markdown parts
Combines all PROJECT_REPORT_PART*.md files into a single formatted Word document
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

def set_cell_border(cell, **kwargs):
    """
    Set cell borders
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Create borders element
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '4')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), '000000')
            tcBorders.append(edge_el)
    tcPr.append(tcBorders)

def add_page_break(doc):
    """Add a page break"""
    doc.add_page_break()

def set_paragraph_format(paragraph, font_name='Times New Roman', font_size=12, 
                         bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         spacing_before=0, spacing_after=0, line_spacing=1.5):
    """Set paragraph formatting"""
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(spacing_before)
    paragraph.paragraph_format.space_after = Pt(spacing_after)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = line_spacing
    
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
        run.italic = italic

def process_markdown_line(doc, line, in_code_block=False, in_table=False):
    """Process a single markdown line and add to document"""
    
    # Skip empty lines in tables
    if in_table and line.strip() == '':
        return in_code_block, in_table
    
    # Handle code blocks
    if line.strip().startswith('```'):
        return not in_code_block, in_table
    
    if in_code_block:
        p = doc.add_paragraph(line, style='Normal')
        set_paragraph_format(p, font_name='Courier New', font_size=10)
        return in_code_block, in_table
    
    # Handle tables
    if line.strip().startswith('|') and '|' in line:
        return in_code_block, True
    elif in_table and not line.strip().startswith('|'):
        return in_code_block, False
    
    # Handle headings
    if line.startswith('# '):
        # Title heading (H1)
        p = doc.add_paragraph(line[2:].strip())
        set_paragraph_format(p, font_size=36, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                           spacing_before=12, spacing_after=12)
    elif line.startswith('## '):
        # Main heading (H2)
        p = doc.add_paragraph(line[3:].strip())
        set_paragraph_format(p, font_size=16, bold=True, spacing_before=12, spacing_after=6)
    elif line.startswith('### '):
        # Sub heading (H3)
        p = doc.add_paragraph(line[4:].strip())
        set_paragraph_format(p, font_size=14, bold=True, italic=True, spacing_before=6, spacing_after=6)
    elif line.startswith('#### '):
        # Minor heading (H4)
        p = doc.add_paragraph(line[5:].strip())
        set_paragraph_format(p, font_size=12, bold=True, spacing_before=6, spacing_after=3)
    
    # Handle horizontal rules
    elif line.strip() == '---':
        p = doc.add_paragraph('_' * 80)
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    # Handle bullet points
    elif line.strip().startswith('- ') or line.strip().startswith('* '):
        text = line.strip()[2:]
        p = doc.add_paragraph(text, style='List Bullet')
        set_paragraph_format(p, font_size=12)
    
    # Handle numbered lists
    elif re.match(r'^\d+\.\s', line.strip()):
        text = re.sub(r'^\d+\.\s', '', line.strip())
        p = doc.add_paragraph(text, style='List Number')
        set_paragraph_format(p, font_size=12)
    
    # Handle bold text
    elif line.strip().startswith('**') and line.strip().endswith('**'):
        p = doc.add_paragraph(line.strip().strip('**'))
        set_paragraph_format(p, font_size=12, bold=True)
    
    # Regular paragraph
    elif line.strip() and not in_table:
        # Process inline formatting
        p = doc.add_paragraph()
        
        # Split by bold markers
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.strip():
                run = p.add_run(part)
        
        set_paragraph_format(p, font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    return in_code_block, in_table

def process_table(doc, table_lines):
    """Process markdown table and add to document"""
    if len(table_lines) < 2:
        return
    
    # Parse table
    rows = []
    for line in table_lines:
        if '|' in line:
            cells = [cell.strip() for cell in line.split('|')]
            cells = [c for c in cells if c]  # Remove empty cells
            if cells and not all(c.startswith('-') for c in cells):  # Skip separator line
                rows.append(cells)
    
    if not rows:
        return
    
    # Create table
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    
    # Fill table
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                cell.text = cell_data
                
                # Format header row
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(11)
                else:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

def add_header_footer(doc, header_text="SmartAd Optimizer", footer_text="PG Department of Computer Applications"):
    """Add header and footer to document"""
    # Add header
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = header_text
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header_para.runs:
        run.font.size = Pt(10)
        run.font.italic = True
    
    # Add footer with page number
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = footer_text + " | Page "
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add page number field
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    for run in footer_para.runs:
        run.font.size = Pt(10)

def set_page_margins(doc):
    """Set page margins"""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1.25)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(0.7)

def generate_report():
    """Main function to generate the report"""
    
    # Create document
    doc = Document()
    
    # Set page margins
    set_page_margins(doc)
    
    # List of report parts in order
    report_parts = [
        'PROJECT_REPORT_PART1.md',
        'PROJECT_REPORT_PART2.md',
        'PROJECT_REPORT_PART3.md',
        'PROJECT_REPORT_PART4.md',
        'PROJECT_REPORT_PART5.md',
        'PROJECT_REPORT_PART6.md',
        'PROJECT_REPORT_PART7.md',
        'PROJECT_REPORT_FINAL.md'
    ]
    
    print("Generating SmartAd Optimizer Project Report...")
    
    # Process each part
    for part_num, part_file in enumerate(report_parts, 1):
        if not os.path.exists(part_file):
            print(f"Warning: {part_file} not found, skipping...")
            continue
        
        print(f"Processing {part_file}... ({part_num}/{len(report_parts)})")
        
        with open(part_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        in_code_block = False
        in_table = False
        table_lines = []
        
        for line in lines:
            line = line.rstrip('\n')
            
            # Handle tables
            if line.strip().startswith('|'):
                in_table = True
                table_lines.append(line)
            elif in_table:
                if line.strip() == '':
                    # End of table
                    process_table(doc, table_lines)
                    table_lines = []
                    in_table = False
                else:
                    table_lines.append(line)
            else:
                # Process regular lines
                in_code_block, in_table = process_markdown_line(doc, line, in_code_block, in_table)
        
        # Process any remaining table
        if table_lines:
            process_table(doc, table_lines)
        
        # Add page break between parts (except for last part)
        if part_num < len(report_parts):
            add_page_break(doc)
    
    # Add header and footer
    add_header_footer(doc)
    
    # Save document
    output_file = 'SmartAd_Optimizer_Project_Report.docx'
    doc.save(output_file)
    
    print(f"\n[SUCCESS] Report generated successfully: {output_file}")
    print(f"\nNext steps:")
    print("1. Open the document in Microsoft Word")
    print("2. Replace placeholders: [YOUR NAME], [YOUR REG NO], [INTERNAL GUIDE NAME]")
    print("3. Add actual screenshots for figures")
    print("4. Update GitHub and Colab links")
    print("5. Review formatting and adjust if needed")
    print("6. Add page numbers starting from Introduction chapter")
    print("7. Print and soft bind as per requirements")

if __name__ == '__main__':
    try:
        generate_report()
    except Exception as e:
        print(f"Error generating report: {e}")
        import traceback
        traceback.print_exc()
